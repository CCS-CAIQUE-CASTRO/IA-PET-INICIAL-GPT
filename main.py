from __future__ import annotations
import os, sys, json
from pathlib import Path
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document  # python-docx
from docx.shared import RGBColor,Pt

# ---------- paths base ----------

def _app_base() -> Path:
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent

env_path = _app_base() / ".env"
load_dotenv(dotenv_path=env_path, override=False)

# modelo Word (coloque esse arquivo ao lado do .exe ou do main.py)
MODELO_WORD = _app_base() / "Pet Inicial modelo para IA.docx"

_client: OpenAI | None = None

def _get_client() -> OpenAI:
    global _client
    if _client is None:
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise RuntimeError(
                "OPENAI_API_KEY não encontrada.\n"
                f"Procurei em: {env_path}\n"
                "Crie um arquivo .env ao lado do executável com:\n"
                "OPENAI_API_KEY=sua_chave_aqui"
            )
        _client = OpenAI(api_key=api_key)
    return _client

# ID do agente salvo
_PROMPT_ID = "pmpt_691610f1f92c8195813434067cd51f490c72f28960d54f47"


# ---------- helpers para JSON e Word ----------

def _extrair_json_puro(texto: str) -> str:
    """
    Remove ```json ... ``` ou ``` ... ``` se o modelo devolver em bloco de código.
    """
    texto = texto.strip()
    if texto.startswith("```"):
        linhas = texto.splitlines()
        # remove a primeira linha ``` ou ```json
        if linhas and linhas[0].startswith("```"):
            linhas = linhas[1:]
        # remove última linha ``` se tiver
        if linhas and linhas[-1].startswith("```"):
            linhas = linhas[:-1]
        texto = "\n".join(linhas).strip()
    return texto

def preencher_modelo_word(
    dados: dict[str, str],
    caminho_modelo: Path | None = None,
) -> Path:
    """
    Abre o modelo Word (com placeholders tipo {NOME_CLIENTE}) e substitui
    pelos valores do dicionário `dados`.
    """
    if caminho_modelo is None:
        caminho_modelo = MODELO_WORD

    if not caminho_modelo.exists():
        raise FileNotFoundError(f"Modelo Word não encontrado: {caminho_modelo}")

    doc = Document(str(caminho_modelo))
    AZUL = RGBColor(0, 0, 255)

    # ---------- PREPARAÇÃO DOS DADOS COM REGRA ESPECIAL PARA VINCULO_COM_TRABALHO ----------
    dados_preparados = {}
    
    # Copia todos os dados, tratando valores None e strings vazias
    for chave, valor in dados.items():
        if valor is None:
            continue
        valor_str = str(valor).strip()
        if valor_str:
            dados_preparados[chave] = valor_str
    
    # REGRA ESPECIAL: VINCULO_COM_TRABALHO
    # Primeiro verifica se existe esse campo nos dados
    if "VINCULO_COM_TRABALHO" in dados_preparados:
        vinculo_valor = dados_preparados["VINCULO_COM_TRABALHO"].upper().strip()
        if vinculo_valor == "SIM":
            dados_preparados["VINCULO_COM_TRABALHO"] = (
                "C)\tDA JUSTIÇA GRATUITA\n\n"
                "Primeiramente, o art. 129, par. único da Lei 8.213/91  garante a isenção quanto a "
                "custas e verbas sucumbenciais nas causas decorrentes de acidente do trabalho.\n\n"
                "Além da isenção garantida pela lei, é importante aludir que a parte Autora não tem "
                "condições de arcar com quaisquer custas, despesas e/ou honorários advocatícios sem "
                "prejuízo do próprio sustento.\n\n"
                "Portanto, requer-se o reconhecimento da isenção conferida pelo art. 129, par. único "
                "da Lei 8.213/91, abstendo a parte autora de qualquer ônus monetário no presente caso "
                "ou, que sejam deferidos os benefícios da justiça gratuita, nos moldes Lei 1.060/50, "
                "bem como, artigo 98 e seguintes do CPC, considerando declaração de hipossuficiência "
                "e provas a ela anexadas."
            )
        else:
            # Se não for "SIM", remove completamente do dicionário
            # Isso faz com que o placeholder {VINCULO_COM_TRABALHO} seja mantido no texto
            # e depois substituído por string vazia no processamento
            dados_preparados["VINCULO_COM_TRABALHO"] = ""

    # ---------- FUNÇÃO PARA PROCESSAR PARÁGRAFOS ----------
    def processar_paragrafo(p):
        full = p.text
        if "{" not in full or "}" not in full:
            return

        segmentos: list[tuple[str, str]] = []
        i = 0
        n = len(full)

        while i < n:
            if full[i] == "{":
                j = full.find("}", i + 1)
                if j == -1:
                    segmentos.append(("fixo", full[i:]))
                    break

                chave = full[i + 1:j]
                # SEMPRE tenta substituir, mesmo que a chave não esteja em dados_preparados
                # Se não estiver, usa string vazia
                valor = dados_preparados.get(chave, "")
                segmentos.append(("valor", valor))
                i = j + 1
            else:
                prox = full.find("{", i)
                if prox == -1:
                    segmentos.append(("fixo", full[i:]))
                    break
                else:
                    segmentos.append(("fixo", full[i:prox]))
                    i = prox

        # Mantém o estilo do parágrafo
        estilo = p.style
        p.clear()
        p.style = estilo

        # Recria os runs
        for tipo, texto in segmentos:
            if not texto:
                continue
            run = p.add_run(texto)
            if tipo == "valor":
                run.font.color.rgb = AZUL
                run.font.name = "Bookman Old Style"
                run.font.size = Pt(12)

    # Processa parágrafos normais
    for p in doc.paragraphs:
        processar_paragrafo(p)

    # Processa tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    processar_paragrafo(p)

    # Gera nome do arquivo de saída
    nome_base = dados_preparados.get("NOME_CLIENTE", "Pet_inicial")
    nome_base = str(nome_base).strip()
    nome_base = "".join(
        c for c in nome_base if c.isalnum() or c in (" ", "_", "-")
    ).strip() or "Pet_inicial"

    saida_path = caminho_modelo.with_name(f"Pet_inicial_{nome_base}.docx")
    doc.save(str(saida_path))
    return saida_path





# ---------- função para até 5 PDFs + geração opcional do Word ----------

def analisar_pdfs(caminhos_pdfs: list[Path], gerar_word: bool = False) -> str:
    """
    Envia até 5 PDFs para o agente e retorna a resposta em texto.
    Se gerar_word=True, tenta interpretar a resposta como JSON e
    preenche o modelo Word.

    Retorno:
        - sempre uma string com a saída da IA;
        - se gerar_word=True, acrescenta no final da string o caminho
          do .docx gerado (ou uma mensagem de erro se não entender o JSON).
    """
    if not caminhos_pdfs:
        return "Erro: nenhum PDF informado."

    # garante no máximo 5
    caminhos_pdfs = caminhos_pdfs[:5]

    # valida arquivos
    pdfs_validos: list[Path] = []
    for p in caminhos_pdfs:
        if not p.exists() or not p.is_file():
            return f"Erro: arquivo não encontrado: {p}"
        if p.suffix.lower() != ".pdf":
            return f"Erro: arquivo não é PDF: {p}"
        pdfs_validos.append(p)

    client = _get_client()

    file_ids: list[str] = []
    try:
        # 1) faz upload de todos os PDFs
        for p in pdfs_validos:
            with open(p, "rb") as f:
                uploaded = client.files.create(file=f, purpose="assistants")
            file_ids.append(uploaded.id)

        # 2) monta o conteúdo da mensagem pro agente
        conteudo = []

        # adiciona cada PDF como input_file
        for fid in file_ids:
            conteudo.append({"type": "input_file", "file_id": fid})

        # 3) chama o agente salvo
        response = client.responses.create(
            prompt={"id": _PROMPT_ID},
            input=[
                {
                    "role": "user",
                    "content": conteudo,
                }
            ],
        )

        saida = (response.output_text or "").strip()
        if not saida:
            saida = "(Sem conteúdo retornado pelo agente.)"

        # 4) se quiser já gerar o Word aqui
        if gerar_word:
            try:
                json_puro = _extrair_json_puro(saida)
                dados = json.loads(json_puro)
                if not isinstance(dados, dict):
                    raise ValueError("JSON não é um objeto/dicionário.")

                caminho_docx = preencher_modelo_word(dados)
                saida += f"\n\n✅ Petição Word gerada em:\n{caminho_docx}"
            except Exception as exc:
                saida += (
                    "\n\n[AVISO] Não foi possível gerar o Word a partir da "
                    f"resposta JSON: {exc!s}"
                )

        return saida

    except Exception as exc:
        return f"Erro ao consultar a IA: {exc!s}"

    finally:
        # 5) tenta deletar os arquivos enviados (boa prática)
        try:
            if file_ids:
                for fid in file_ids:
                    try:
                        client.files.delete(fid)
                    except Exception:
                        pass
        except Exception:
            pass


if __name__ == "__main__":
    # exemplo com 5 PDFs
    pdfs = [
        Path("05. CAT - 12.06.2019 - Trajeto.pdf"),
        Path("06. Laudo Médico - Pericial.pdf"),
        Path("07. Decl. de Benefício.pdf"),
        Path("09.1. DECLARAÇÃO + PROCURAÇÃO - ASSINADOS!.pdf"),
        Path("Roteiro de Visita.pdf"),
    ]
    print("✅ Enviando PDFs:", pdfs)
    texto = analisar_pdfs(pdfs, gerar_word=True)
    print("\n📄 Resultado da análise:\n")
    print(texto)
    
