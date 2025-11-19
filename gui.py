import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Optional, Callable, List


from PySide6.QtCore import Qt, QMimeData, QThread, Signal, QSize
from PySide6.QtGui import QDragEnterEvent, QDropEvent, QAction
from PySide6.QtWidgets import (
    QApplication,
    QMainWindow,
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QPushButton,
    QLabel,
    QTextEdit,
    QFileDialog,
    QProgressBar,
    QMessageBox,
    QFrame,
    QStyle,
    QSpacerItem,
    QSizePolicy,
)

from docx import Document  # python-docx


# ==========================
# Caminho do modelo Word padrão
# ==========================
def _app_base() -> Path:
    """
    Retorna a pasta base da aplicação:
    - Se empacotado (PyInstaller onefile), usa a pasta do executável.
    - Caso contrário, usa a pasta do próprio arquivo .py.
    """
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


TEMPLATE_DOCX = _app_base() / "Pet Inicial modelo para IA.docx"


# ==========================
# “Agente de IA” (implementação real)
# ==========================
def chamar_agente_ia_pdf(caminho_pdf: Path) -> str:
    """
    Encaminha para a função real do seu agente em main.py.
    Mantém a mesma assinatura esperada pelo IAWorker.
    """
    from main import analisar_pdf
    return analisar_pdf(caminho_pdf)


# ==========================
# Worker em thread (não travar UI)
# ==========================
@dataclass
class JobParams:
    arquivos: List[Path]
    func: Callable[[Path], str]



class IAWorker(QThread):
    progressed = Signal(int)
    finished_ok = Signal(str)
    failed = Signal(str)

    def __init__(self, params: JobParams):
        super().__init__()
        self.params = params
        self._cancel = False

    def cancel(self) -> None:
        self._cancel = True

    def run(self) -> None:
        try:
            # Fase inicial de "loading"
            for i in range(0, 40):
                if self._cancel:
                    self.failed.emit("Execução cancelada.")
                    return
                self.progressed.emit(i + 1)
                time.sleep(0.015)

            if self._cancel:
                self.failed.emit("Execução cancelada.")
                return

            # Processa TODOS os PDFs
            textos = []
            for pdf_path in self.params.arquivos:
                if self._cancel:
                    self.failed.emit("Execução cancelada.")
                    return
                resp = self.params.func(pdf_path)
                textos.append(f"=== {pdf_path.name} ===\n{resp}")

            # Fase final de "loading"
            for i in range(40, 100):
                if self._cancel:
                    self.failed.emit("Execução cancelada.")
                    return
                self.progressed.emit(i + 1)
                time.sleep(0.006)

            # Junta tudo em um único texto
            resultado_final = "\n\n\n".join(textos)
            self.finished_ok.emit(resultado_final)

        except Exception as exc:
            self.failed.emit(f"Erro ao consultar a IA: {exc!s}")


# ==========================
# Widget de dropzone (PDF-only)
# ==========================
class DropZone(QFrame):
    file_selected = Signal(Path)

    def __init__(self) -> None:
        super().__init__()
        self.setAcceptDrops(True)
        self.setObjectName("DropZone")
        self.setFrameShape(QFrame.StyledPanel)
        self.setMinimumHeight(120)

        lay = QVBoxLayout(self)
        self.icon_label = QLabel()
        self.icon_label.setAlignment(Qt.AlignCenter)
        self.icon_label.setPixmap(
            self.style().standardIcon(QStyle.SP_DialogOpenButton).pixmap(QSize(48, 48))
        )
        self.title = QLabel("Arraste um PDF aqui")
        self.title.setAlignment(Qt.AlignCenter)
        self.subtitle = QLabel('ou clique em "Selecionar PDF"')
        self.subtitle.setAlignment(Qt.AlignCenter)

        lay.addWidget(self.icon_label)
        lay.addWidget(self.title)
        lay.addWidget(self.subtitle)

    def dragEnterEvent(self, event: QDragEnterEvent) -> None:
        md: QMimeData = event.mimeData()
        if md.hasUrls():
            for u in md.urls():
                if u.toLocalFile().lower().endswith(".pdf"):
                    event.acceptProposedAction()
                    return
        event.ignore()

    def dropEvent(self, event: QDropEvent) -> None:
        urls = event.mimeData().urls()
        if not urls:
            return
        local = urls[0].toLocalFile()
        if local and local.lower().endswith(".pdf"):
            self.file_selected.emit(Path(local))
        else:
            QMessageBox.warning(self, "Arquivo inválido", "Envie apenas arquivos PDF.")


# ==========================
# Janela Principal
# ==========================
class MainWindow(QMainWindow):
    def __init__(self) -> None:
        super().__init__()
        self.setWindowTitle("Assistente IA — PDF")
        self.setMinimumSize(900, 580)
        self.setWindowIcon(self.style().standardIcon(QStyle.SP_ComputerIcon))

        self._arquivos: List[Path] = []

        self._worker: Optional[IAWorker] = None

        self._build_ui()
        self._apply_style(dark=True)

    # ---------- UI ----------
    def _build_ui(self) -> None:
        central = QWidget()
        root = QVBoxLayout(central)
        root.setContentsMargins(16, 16, 16, 16)
        root.setSpacing(12)

        # Barra de ações (arquivo + tema)
        bar = QHBoxLayout()
        self.btn_sel = QPushButton("Selecionar PDF…")
        self.btn_sel.clicked.connect(self.on_select_file)
        self.lbl_arquivo = QLabel("Nenhum PDF selecionado")
        self.lbl_arquivo.setObjectName("ArquivoLabel")
        bar.addWidget(self.btn_sel)
        bar.addSpacing(8)
        bar.addWidget(self.lbl_arquivo, 1)

        self.btn_theme = QPushButton("Tema: Escuro")
        self.btn_theme.setCheckable(True)
        self.btn_theme.setChecked(True)
        self.btn_theme.clicked.connect(self.on_toggle_theme)
        bar.addSpacing(12)
        bar.addWidget(self.btn_theme)

        root.addLayout(bar)

        # Dropzone
        self.drop = DropZone()
        self.drop.file_selected.connect(self.set_file)
        root.addWidget(self.drop)

        # Ações: executar/cancelar + progresso
        actions = QHBoxLayout()
        self.btn_run = QPushButton("Perguntar à IA")
        self.btn_run.setDefault(True)
        self.btn_run.clicked.connect(self.on_run)

        self.btn_cancel = QPushButton("Cancelar")
        self.btn_cancel.setEnabled(False)
        self.btn_cancel.clicked.connect(self.on_cancel)

        actions.addWidget(self.btn_run)
        actions.addWidget(self.btn_cancel)
        actions.addItem(QSpacerItem(20, 20, QSizePolicy.Expanding, QSizePolicy.Minimum))

        self.progress = QProgressBar()
        self.progress.setMinimum(0)
        self.progress.setMaximum(100)
        self.progress.setValue(0)
        self.progress.setTextVisible(True)
        actions.addWidget(self.progress)

        root.addLayout(actions)

        # Saída da IA
        out_lbl = QLabel("Resposta da IA:")
        self.out = QTextEdit()
        self.out.setReadOnly(True)
        self.out.setPlaceholderText("A resposta aparecerá aqui…")
        root.addWidget(out_lbl)
        root.addWidget(self.out, 1)

        # Rodapé: copiar / limpar / gerar Word
        footer = QHBoxLayout()
        self.btn_copy = QPushButton("Copiar resposta")
        self.btn_copy.clicked.connect(self.on_copy)

        self.btn_clear = QPushButton("Limpar")
        self.btn_clear.clicked.connect(self.on_clear)

        self.btn_word = QPushButton("Gerar Word")
        self.btn_word.clicked.connect(self.on_generate_word)

        footer.addWidget(self.btn_copy)
        footer.addWidget(self.btn_clear)
        footer.addWidget(self.btn_word)
        footer.addItem(QSpacerItem(20, 20, QSizePolicy.Expanding, QSizePolicy.Minimum))

        root.addLayout(footer)

        self.setCentralWidget(central)

        # Menu “Sobre”
        about_act = QAction("Sobre", self)
        about_act.triggered.connect(self.on_about)
        self.menuBar().addMenu("Ajuda").addAction(about_act)

    # ---------- Estilo ----------
    def _apply_style(self, *, dark: bool) -> None:
        base_fg = "#EAECEF" if dark else "#1F2328"
        base_bg = "#0D1117" if dark else "#FFFFFF"
        card_bg = "#161B22" if dark else "#F6F8FA"
        accent = "#2F81F7" if dark else "#0969DA"
        subtle = "#8B949E" if dark else "#57606A"

        self.setStyleSheet(f"""
            QMainWindow {{
                background: {base_bg};
                color: {base_fg};
            }}
            QLabel {{
                color: {base_fg};
                font-size: 14px;
            }}
            #ArquivoLabel {{
                color: {subtle};
            }}
            QFrame#DropZone {{
                background: {card_bg};
                border: 2px dashed {subtle};
                border-radius: 16px;
            }}
            QPushButton {{
                background: {card_bg};
                border: 1px solid rgba(127,127,127,0.25);
                border-radius: 10px;
                padding: 8px 12px;
                color: {base_fg};
            }}
            QPushButton:hover {{
                border-color: {accent};
            }}
            QPushButton:disabled {{
                opacity: .5;
            }}
            QTextEdit {{
                background: {card_bg};
                border: 1px solid rgba(127,127,127,0.25);
                border-radius: 10px;
                padding: 8px 10px;
                color: {base_fg};
                selection-background-color: {accent};
            }}
            QProgressBar {{
                background: {card_bg};
                border: 1px solid rgba(127,127,127,0.25);
                border-radius: 10px;
                text-align: center;
                color: {base_fg};
                min-width: 220px;
                padding: 2px;
            }}
            QProgressBar::chunk {{
                background-color: {accent};
                border-radius: 8px;
            }}
            QMenuBar, QMenu {{
                background: {card_bg};
                color: {base_fg};
            }}
            QMenu::item:selected {{
                background: {accent};
                color: white;
            }}
        """)

    # ---------- Lógica ----------
    def on_select_file(self) -> None:
        paths, _ = QFileDialog.getOpenFileNames(
        self,
        "Selecione um ou mais PDFs",
        filter="PDF (*.pdf)"
    )
        
        if paths:
            arquivos = [Path(p) for p in paths]
            self.set_files(arquivos)


    def set_files(self, arquivos: List[Path]) -> None:
        # Garante que todos são PDFs
        validos = [p for p in arquivos if p.suffix.lower() == ".pdf"]
        if not validos:
            QMessageBox.warning(self, "Arquivo inválido", "Envie apenas arquivos PDF.")
            return

        self._arquivos = validos

        if len(validos) == 1:
            self.lbl_arquivo.setText(str(validos[0]))
        else:
            nomes = ", ".join(p.name for p in validos)
            self.lbl_arquivo.setText(f"{len(validos)} PDFs selecionados: {nomes}")

    def set_file(self, p: Path) -> None:
        # Mantém compatibilidade com o DropZone (um arquivo só)
        if p.suffix.lower() != ".pdf":
            QMessageBox.warning(self, "Arquivo inválido", "Envie apenas arquivos PDF.")
            return
        self.set_files([p])


    def on_toggle_theme(self) -> None:
        dark = self.btn_theme.isChecked()
        self.btn_theme.setText("Tema: Escuro" if dark else "Tema: Claro")
        self._apply_style(dark=dark)

    def on_run(self) -> None:
        if self._worker and self._worker.isRunning():
            return
        if not self._arquivos:
            QMessageBox.warning(
                self,
                "PDF obrigatório",
                "Selecione ou arraste pelo menos um arquivo PDF antes de executar.",
            )
            return

        params = JobParams(arquivos=self._arquivos, func=chamar_agente_ia_pdf)
        self._worker = IAWorker(params)
        self._worker.progressed.connect(self.progress.setValue)
        self._worker.finished_ok.connect(self._on_finished_ok)
        self._worker.failed.connect(self._on_failed)

        self.btn_run.setEnabled(False)
        self.btn_cancel.setEnabled(True)
        self.progress.setValue(0)
        self.out.clear()
        self._worker.start()


    def on_cancel(self) -> None:
        if self._worker and self._worker.isRunning():
            self._worker.cancel()

    def _on_finished_ok(self, text: str) -> None:
        self.btn_run.setEnabled(True)
        self.btn_cancel.setEnabled(False)
        self.progress.setValue(100)
        self.out.setPlainText(text)

    def _on_failed(self, msg: str) -> None:
        self.btn_run.setEnabled(True)
        self.btn_cancel.setEnabled(False)
        self.progress.setValue(0)
        QMessageBox.critical(self, "Erro", msg)

    def on_copy(self) -> None:
        txt = self.out.toPlainText().strip()
        if not txt:
            QMessageBox.information(self, "Copiar", "Não há conteúdo para copiar.")
            return
        QApplication.clipboard().setText(txt)
        QMessageBox.information(self, "Copiar", "Resposta copiada!")

    def on_clear(self) -> None:
        self.out.clear()

    def on_generate_word(self) -> None:
        """
        Gera um Word com base em um modelo padrão (TEMPLATE_DOCX, se existir)
        contendo:
        - Arquivo PDF analisado
        - Resposta da IA
        """
        if not self._arquivo:
            QMessageBox.warning(
                self,
                "PDF obrigatório",
                "Selecione ou arraste um arquivo PDF antes de gerar o Word.",
            )
            return

        # Caminho para salvar o Word
        save_path, _ = QFileDialog.getSaveFileName(
            self,
            "Salvar Word",
            "Pet Inicial IA.docx",
            "Documentos Word (*.docx)",
        )
        if not save_path:
            return

        resposta = self.out.toPlainText().strip()
        if not resposta:
            resposta = "Nenhuma resposta da IA foi gerada ainda."

        # Abre o modelo, se existir; senão cria um documento novo
        if TEMPLATE_DOCX.exists():
            doc = Document(TEMPLATE_DOCX)
        else:
            doc = Document()
            doc.add_heading("Petição Inicial - IA", level=1)

        # Adiciona informações do PDF
        doc.add_paragraph()
        doc.add_heading("Arquivo analisado", level=2)
        doc.add_paragraph(str(self._arquivo))

        # Adiciona a resposta da IA
        doc.add_paragraph()
        doc.add_heading("Resposta da IA", level=2)
        for linha in resposta.splitlines():
            if linha.strip():
                doc.add_paragraph(linha)

        try:
            doc.save(save_path)
        except Exception as e:
            QMessageBox.critical(
                self,
                "Erro ao salvar",
                f"Não foi possível salvar o arquivo Word:\n{e}",
            )
            return

        QMessageBox.information(
            self,
            "Word gerado",
            "Documento Word gerado com sucesso!",
        )

    def on_about(self) -> None:
        QMessageBox.information(
            self,
            "Sobre",
            "Assistente IA — PDF\n"
            "Envie um PDF, obtenha a resposta do agente de IA e gere um Word padrão.",
        )


if __name__ == "__main__":
    app = QApplication(sys.argv)
    win = MainWindow()
    win.show()
    sys.exit(app.exec())
